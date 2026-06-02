import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Sets background color of a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets internal padding/margins of a cell in twentieths of a point (dxa)."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_report():
    doc = docx.Document()
    
    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Color palette
    PRIMARY_COLOR = RGBColor(31, 78, 121)     # Steel Blue
    SECONDARY_COLOR = RGBColor(112, 128, 144) # Slate Grey
    TEXT_COLOR = RGBColor(51, 51, 51)         # Dark Grey
    ACCENT_WARN = RGBColor(192, 57, 43)       # Rust Red
    
    # Base Style modifications
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = TEXT_COLOR
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(6)
    
    # Document Title
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(12)
    title.paragraph_format.space_before = Pt(6)
    run_title = title.add_run("Day 1 Data Ingestion & Quality Summary")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = PRIMARY_COLOR
    
    # Subtitle/Metadata
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(24)
    run_meta = meta.add_run("Mutual Fund Analytics Capstone Project | Day 1 Work Review (Real Datasets)")
    run_meta.font.size = Pt(10)
    run_meta.font.italic = True
    run_meta.font.color.rgb = SECONDARY_COLOR
    
    # Intro
    p = doc.add_paragraph("Hey,")
    p = doc.add_paragraph(
        "Here is the complete summary of the work I did today for the Mutual Fund Analytics setup. "
        "I've got the project folder structure organized, the local environment set up, all the raw "
        "datasets ingested, and the live API connection working."
    )
    p = doc.add_paragraph(
        "I did a deep dive into the 10 real CSV files to check for quality issues. The datasets are "
        "practically pristine with no duplicates and excellent relational consistency. There is only "
        "one minor mathematically expected missing field which we'll address in the cleaning phase."
    )
    
    # Section 1
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)
    run = h1.add_run("1. Project Directory & Environment Layout")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR
    
    p = doc.add_paragraph(
        "I verified the project folder structure inside the directory: "
    )
    p.add_run("C:\\Users\\jibum\\OneDrive\\Desktop\\Bluestock Internship").bold = True
    
    p = doc.add_paragraph("Here is the current layout and where we stand:")
    
    # Bullet points
    bp1 = doc.add_paragraph(style='List Bullet')
    bp1.add_run("data/raw/").bold = True
    bp1.add_run(" - Active. Contains the 10 prefixed local CSV datasets, plus the raw historical CSVs downloaded from the API.")
    
    bp2 = doc.add_paragraph(style='List Bullet')
    bp2.add_run("data/processed/").bold = True
    bp2.add_run(" - Empty. An empty staging folder. Deduplicated, clean, and merged datasets will be saved here during the Day 2 cleaning phase.")
    
    bp3 = doc.add_paragraph(style='List Bullet')
    bp3.add_run("notebooks/").bold = True
    bp3.add_run(" - Placeholder. Set up and ready for Jupyter notebooks for explorative analysis (EDA).")
    
    bp4 = doc.add_paragraph(style='List Bullet')
    bp4.add_run("sql/").bold = True
    bp4.add_run(" - Placeholder. Ready to hold database table definitions, schemas, and staging queries.")
    
    bp5 = doc.add_paragraph(style='List Bullet')
    bp5.add_run("dashboard/").bold = True
    bp5.add_run(" - Staging. Set up for Power BI visual assets and dashboard files.")
    
    bp6 = doc.add_paragraph(style='List Bullet')
    bp6.add_run("reports/").bold = True
    bp6.add_run(" - Active. Where documentation, summaries, and data quality reports are stored.")
    
    # Quick Note
    p_note = doc.add_paragraph()
    p_note.paragraph_format.space_before = Pt(8)
    p_note.add_run("Quick Note on Environment: ").bold = True
    p_note.add_run(
        "I initialized the local Git repository and set up the remote pointing to our GitHub repository. "
        "All changes are committed cleanly under \"Day 1: Data ingestion complete\"."
    )
    p_note = doc.add_paragraph(
        "I also created requirements.txt with all dependencies for the analytical stack (Pandas, NumPy, "
        "Matplotlib, Seaborn, Plotly, SQLAlchemy, Requests, SciPy, and Jupyter)."
    )
    
    # Section 2
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)
    run = h2.add_run("2. Ingested Datasets & Structural Profiles")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR
    
    p = doc.add_paragraph(
        "I updated the ETL script (data_ingestion.py) to systematically load and inspect all 10 real source CSV files using Pandas. "
        "Here is the exact profile of the data as it stands:"
    )
    
    # Table Creation
    headers = ["Dataset File Name", "Shape", "Target Primary Key", "Main Column Dtypes", "Quality Status"]
    data = [
        ["01_fund_master.csv", "40 x 15", "amfi_code", "int64, object (string)", "Clean"],
        ["02_nav_history.csv", "46,000 x 3", "amfi_code + date", "int64, object (string), float64", "Clean"],
        ["03_aum_by_fund_house.csv", "90 x 5", "date + fund_house", "object, float64, int64", "Clean"],
        ["04_monthly_sip_inflows.csv", "48 x 6", "month", "object, int64, float64", "Minor Anomaly"],
        ["05_category_inflows.csv", "144 x 3", "month + category", "object, float64", "Clean"],
        ["06_industry_folio_count.csv", "21 x 6", "month", "object, float64", "Clean"],
        ["07_scheme_performance.csv", "40 x 19", "amfi_code", "int64, object, float64", "Clean"],
        ["08_investor_transactions.csv", "32,778 x 13", "investor_id + date", "object, int64, float64", "Clean"],
        ["09_portfolio_holdings.csv", "322 x 8", "amfi_code + stock", "int64, object, float64", "Clean"],
        ["10_benchmark_indices.csv", "8,050 x 3", "date + index_name", "object, float64", "Clean"]
    ]
    
    table = doc.add_table(rows=1 + len(data), cols=5)
    table.style = 'Light Shading Accent 1'
    
    # Headers
    hdr_cells = table.rows[0].cells
    for idx, text in enumerate(headers):
        hdr_cells[idx].text = text
        set_cell_background(hdr_cells[idx], "1F4E79")
        set_cell_margins(hdr_cells[idx], 120, 120, 150, 150)
        run = hdr_cells[idx].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
        
    # Data Rows
    for r_idx, row_data in enumerate(data):
        row_cells = table.rows[r_idx + 1].cells
        bg_color = "F2F5F8" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, text in enumerate(row_data):
            row_cells[c_idx].text = text
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], 80, 80, 150, 150)
            run = row_cells[c_idx].paragraphs[0].runs[0]
            run.font.size = Pt(9.5)
            if "Minor Anomaly" in text:
                run.font.bold = True
                run.font.color.rgb = ACCENT_WARN
            elif "Clean" in text:
                run.font.color.rgb = RGBColor(46, 125, 50) # Green
                
    # Section 3
    h3 = doc.add_paragraph()
    h3.paragraph_format.space_before = Pt(24)
    h3.paragraph_format.space_after = Pt(6)
    run = h3.add_run("3. Data Quality & Anomalies Breakdown")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR
    
    p = doc.add_paragraph(
        "While profiling the real-world files, the datasets were found to be very high-quality. Here is the anomalies breakdown:"
    )
    
    # Sub-heading 3.A
    h3a = doc.add_paragraph()
    h3a.paragraph_format.space_before = Pt(8)
    h3a.paragraph_format.space_after = Pt(4)
    run = h3a.add_run("A. Missing (Null) Values")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = SECONDARY_COLOR
    
    bp1 = doc.add_paragraph(style='List Bullet')
    bp1.add_run("04_monthly_sip_inflows.csv: ").bold = True
    bp1.add_run("Contains ")
    bp1.add_run("12 null values").bold = True
    bp1.add_run(" in the yoy_growth_pct column. This is a mathematically expected anomaly because YoY growth calculations require previous year's baseline data which is unavailable for 2022.")
    
    # Sub-heading 3.B
    h3b = doc.add_paragraph()
    h3b.paragraph_format.space_before = Pt(12)
    h3b.paragraph_format.space_after = Pt(4)
    run = h3b.add_run("B. Duplicate Records")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = SECONDARY_COLOR
    
    bp1 = doc.add_paragraph(style='List Bullet')
    bp1.add_run("All CSV Files: ").bold = True
    bp1.add_run("No duplicate rows were found in any of the 10 real-world datasets. Relational key boundaries are completely unique.")
    
    # Sub-heading 3.C
    h3c = doc.add_paragraph()
    h3c.paragraph_format.space_before = Pt(12)
    h3c.paragraph_format.space_after = Pt(4)
    run = h3c.add_run("C. Portfolio Holdings representation")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = SECONDARY_COLOR
    
    bp1 = doc.add_paragraph(style='List Bullet')
    bp1.add_run("09_portfolio_holdings.csv: ").bold = True
    bp1.add_run("Only 34 unique AMFI codes are present (out of the 40). This is accurate as portfolio equity weights are only tracked for equity and index schemes, and naturally exclude the 6 debt/liquid schemes.")
    
    # Section 4
    h4 = doc.add_paragraph()
    h4.paragraph_format.space_before = Pt(18)
    h4.paragraph_format.space_after = Pt(6)
    run = h4.add_run("4. AMFI Code Integrity & Cross-Validation")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR
    
    p = doc.add_paragraph(
        "The AMFI Scheme Code is our natural key to join static metadata with prices or transactions. "
        "I ran a relational consistency check between 01_fund_master.csv and 02_nav_history.csv:"
    )
    
    bp1 = doc.add_paragraph(style='List Bullet')
    bp1.add_run("Unique Master Codes: ").bold = True
    bp1.add_run("40 unique scheme codes detected.")
    
    bp2 = doc.add_paragraph(style='List Bullet')
    bp2.add_run("Unique NAV Codes: ").bold = True
    bp2.add_run("40 unique scheme codes detected.")
    
    # Success Callout Box using Table
    callout = doc.add_table(rows=1, cols=1)
    callout.style = 'Table Grid'
    cell = callout.rows[0].cells[0]
    set_cell_background(cell, "E8F8F5") # Soft Teal background
    set_cell_margins(cell, 150, 150, 200, 200)
    
    p_c = cell.paragraphs[0]
    p_c.paragraph_format.space_after = Pt(0)
    run_c_title = p_c.add_run("SUCCESS ON RELATIONSHIPS:\n")
    run_c_title.font.bold = True
    run_c_title.font.color.rgb = RGBColor(46, 125, 50)
    run_c_title.font.size = Pt(10.5)
    
    run_c_text = p_c.add_run(
        "Relational integrity check shows 100% SUCCESS. Every single AMFI code listed in the master file has "
        "associated historical prices in the daily NAV history, and there are no orphan codes. This ensures a clean star schema."
    )
    run_c_text.font.size = Pt(10)
    
    # Spacer
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(8)
    
    # Section 5
    h5 = doc.add_paragraph()
    h5.paragraph_format.space_before = Pt(18)
    h5.paragraph_format.space_after = Pt(6)
    run = h5.add_run("5. Live Ingestion Results from API (mfapi.in)")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR
    
    p = doc.add_paragraph(
        "I ran the API tool (live_nav_fetch.py) to connect to the open API at https://api.mfapi.in and grab live NAV records. "
        "The script succeeded with 100% retrieval rate and saved CSV files under data/raw/:"
    )
    
    bp1 = doc.add_paragraph(style='List Bullet')
    bp1.add_run("HDFC Top 100 Direct (125497)").bold = True
    bp1.add_run(" -> hdfc_top_100_direct_nav.csv (3,091 daily rows)")
    
    bp2 = doc.add_paragraph(style='List Bullet')
    bp2.add_run("SBI Bluechip (119551)").bold = True
    bp2.add_run(" -> sbi_bluechip_nav.csv (3,236 daily rows)")
    
    bp3 = doc.add_paragraph(style='List Bullet')
    bp3.add_run("ICICI Prudential Bluechip (120503)").bold = True
    bp3.add_run(" -> icici_bluechip_nav.csv (3,307 daily rows)")
    
    bp4 = doc.add_paragraph(style='List Bullet')
    bp4.add_run("Nippon India Large Cap (118632)").bold = True
    bp4.add_run(" -> nippon_large_cap_nav.csv (3,298 daily rows)")
    
    bp5 = doc.add_paragraph(style='List Bullet')
    bp5.add_run("Axis Bluechip (119092)").bold = True
    bp5.add_run(" -> axis_bluechip_nav.csv (3,565 daily rows)")
    
    bp6 = doc.add_paragraph(style='List Bullet')
    bp6.add_run("Kotak Bluechip (120841)").bold = True
    bp6.add_run(" -> kotak_bluechip_nav.csv (3,301 daily rows)")
    
    # Section 6
    h6 = doc.add_paragraph()
    h6.paragraph_format.space_before = Pt(18)
    h6.paragraph_format.space_after = Pt(6)
    run = h6.add_run("6. What's Next (Day 2 Plans)")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR
    
    p = doc.add_paragraph(
        "To get these files ready for our database and dashboards, my plan for Day 2 is to:"
    )
    
    bp1 = doc.add_paragraph(style='List Bullet')
    bp1.add_run("Harmonize Dates: ").bold = True
    bp1.add_run("Ensure all date fields across datasets are parsed into standard YYYY-MM-DD format.")
    
    bp2 = doc.add_paragraph(style='List Bullet')
    bp2.add_run("SQLite Relational Design: ").bold = True
    bp2.add_run("Write DDL statements in sql/schema.sql to set up a normalized star schema in SQLite (bluestock_mf.db).")
    
    bp3 = doc.add_paragraph(style='List Bullet')
    bp3.add_run("Load Data via SQLAlchemy: ").bold = True
    bp3.add_run("Implement an automated load step in Python to populate all 10 datasets into the SQLite database.")
    
    bp4 = doc.add_paragraph(style='List Bullet')
    bp4.add_run("Analytical SQL Verification: ").bold = True
    bp4.add_run("Run the 10 core analytics queries to verify relations and calculations, saving SQL queries in sql/queries.sql.")
    
    p = doc.add_paragraph(
        "Overall, the setup and ingestion went really well today. The provided datasets are very clean and robust. "
        "Let me know if you have any questions or want me to proceed with Day 2 database design!"
    )
    
    # Save Document
    output_path = os.path.join("C:\\Users\\jibum\\OneDrive\\Desktop\\Bluestock Internship\\reports", "day1_data_quality_summary.docx")
    doc.save(output_path)
    print(f"Successfully generated styled Word Document at: {output_path}")

if __name__ == "__main__":
    create_report()

